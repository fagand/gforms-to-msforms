"""Builds .docx files from a Quiz.

Two separate documents are produced — this was a real, confirmed bug fix: an earlier
version appended the Answer Key to the *same* file handed to Quick Import. Quick
Import doesn't stop parsing at the end of the real questions; it kept reading into
that section, absorbed most of it into the last question's own text, and spawned a
garbage extra question out of the leftovers. There is no way to mark trailing content
"ignore this" inside a single Quick-Import document, so the only fix is two files:

- `build_quiz_docx`: the file to actually feed into Quick Import. Only questions and
  choices — nothing else — shaped per docs/MS_FORMS_QUICK_IMPORT.md:
  - Questions are plain paragraphs "<n>. <text>" (numeral + period, no heading style
    required).
  - Choices are plain paragraphs "<LETTER>. <text>" (capital letter + period) — never
    a Word bullet-list style, which is reported to fail Quick Import outright.
  - A blank paragraph separates each question block.
  - No tables, no images, no equations anywhere in the question body.
  - An undocumented-but-tested "ANSWER: <Letter>" / "POINT: <n>" pair is emitted under
    single-choice graded questions, since it costs nothing if Microsoft ignores it.
    No equivalent exists for open-text questions — Microsoft has no import mechanism
    for open-text correct answers at all, so those are always a manual step.
- `build_answer_key_docx`: a separate, human-readable reference document (never fed
  into Quick Import) listing correct answers, points, required flags, and anything
  that needed a judgement call during conversion.
"""
from __future__ import annotations

import re

from docx import Document
from docx.shared import Pt

from .models import Question, QuestionType, Quiz


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().lower()


def _letter(index: int) -> str:
    if 0 <= index < 26:
        return chr(ord("A") + index)
    return str(index + 1)


def _match_option_index(answer: str, options: list[str]) -> int | None:
    target = _normalize(answer)
    for i, option in enumerate(options):
        if _normalize(option) == target:
            return i
    return None


def build_quiz_docx(quiz: Quiz) -> tuple[Document, list[str]]:
    """Returns the Quick-Import-ready document, plus any conversion notes gathered
    while building it (for the separate answer-key document — see module docstring)."""
    doc = Document()
    doc.core_properties.title = quiz.title

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_paragraph(quiz.title, style="Title")
    if quiz.description:
        doc.add_paragraph(quiz.description)
    doc.add_paragraph()

    conversion_notes: list[str] = [f"General: {w}" for w in quiz.warnings]

    for n, question in enumerate(quiz.questions, start=1):
        doc.add_paragraph(f"{n}. {question.text}")

        if question.question_type in (QuestionType.SINGLE_CHOICE, QuestionType.MULTI_CHOICE) and question.options:
            for i, option in enumerate(question.options):
                doc.add_paragraph(f"{_letter(i)}. {option}")

            if question.question_type == QuestionType.SINGLE_CHOICE and question.correct_answers:
                match_index = _match_option_index(question.correct_answers[0], question.options)
                if match_index is not None:
                    doc.add_paragraph(f"ANSWER: {_letter(match_index)}")
                    if question.points is not None:
                        doc.add_paragraph(f"POINT: {question.points}")
                else:
                    conversion_notes.append(
                        f"Q{n}: correct answer text did not match any option exactly; "
                        f"automatic answer marking was skipped for this question."
                    )

        doc.add_paragraph()

        for w in question.warnings:
            conversion_notes.append(f"Q{n}: {w}")

    return doc, conversion_notes


def build_answer_key_docx(quiz: Quiz, conversion_notes: list[str]) -> Document | None:
    """A separate reference document — never fed into Quick Import — listing correct
    answers, points, and required flags per question, plus any conversion notes.
    Returns None if there is nothing worth reporting."""
    answer_key_entries: list[tuple[int, Question]] = [
        (n, q)
        for n, q in enumerate(quiz.questions, start=1)
        if q.correct_answers or q.points is not None or q.required
    ]

    if not answer_key_entries and not conversion_notes:
        return None

    doc = Document()
    doc.core_properties.title = f"{quiz.title} — Answer Key"

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_paragraph(f"{quiz.title} — Answer Key", style="Title")
    note = doc.add_paragraph()
    note_run = note.add_run(
        "This document is for your reference only — do not import it into Microsoft "
        "Forms. Import the other file instead. Microsoft Forms Quick Import does not "
        "guarantee that correct answers, points, or required settings carry over "
        "automatically for any question type — use this list to finish setting them "
        "in Microsoft Forms after import (select each question, then \"Add correct "
        "answers and point values\"). Open-text (typed-answer) questions in "
        "particular always need this done manually; Microsoft has no import "
        "mechanism for open-text correct answers at all."
    )
    note_run.italic = True
    doc.add_paragraph()

    if answer_key_entries:
        doc.add_paragraph("Answer Key", style="Heading 1")
        for n, question in answer_key_entries:
            heading = doc.add_paragraph()
            heading.add_run(f"Q{n}: {question.text}").bold = True

            if question.correct_answers:
                label = "Correct answer(s)" if question.question_type != QuestionType.MULTI_CHOICE else "Correct answer(s) (select all)"
                doc.add_paragraph(f"{label}: {'; '.join(question.correct_answers)}")
            else:
                doc.add_paragraph("Correct answer: not graded in the original form.")

            if question.points is not None:
                doc.add_paragraph(f"Points: {question.points}")
            if question.required:
                doc.add_paragraph("Required: Yes (set this manually — Quick Import cannot set it).")
            if question.question_type == QuestionType.MULTI_CHOICE and question.correct_answers:
                doc.add_paragraph(
                    "Note: enable \"Multiple answers\" on this question after import, "
                    "then tick every correct option listed above."
                )
            doc.add_paragraph()

    if conversion_notes:
        doc.add_paragraph("Conversion Notes", style="Heading 1")
        for note_text in conversion_notes:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(note_text)

    return doc
