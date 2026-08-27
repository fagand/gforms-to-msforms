import io
import zipfile
from pathlib import Path

from docx import Document

from app.converter.pipeline import convert_zip

FIXTURE_ZIP = Path(__file__).parent / "fixtures" / "sample_quiz.zip"


def test_convert_real_sample_end_to_end():
    data = FIXTURE_ZIP.read_bytes()
    result = convert_zip("DR1 Storing Whole Numbers Quiz.zip", data)

    assert result.success is True
    assert result.error is None
    assert result.docx_filename == "DR1 Storing Whole Numbers Quiz.docx"
    assert result.warnings == []

    doc = Document(io.BytesIO(result.docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "1. What number system do computers use to store data?" in text
    assert "ANSWER: C" in text
    # The Answer Key must NOT be in the Quick-Import file itself — a real Microsoft
    # Forms import confirmed it keeps parsing past the real questions, absorbing this
    # section into the last question and spawning a bogus extra question from it.
    assert "Answer Key" not in text

    assert result.answer_key_bytes is not None
    ak_doc = Document(io.BytesIO(result.answer_key_bytes))
    ak_text = "\n".join(p.text for p in ak_doc.paragraphs)
    assert "Answer Key" in ak_text
    assert "Q1: What number system do computers use to store data?" in ak_text
    assert "do not import it into Microsoft Forms" in ak_text


def test_quiz_docx_has_no_trailing_content_after_last_question():
    """Regression test for the real Microsoft Forms import bug: the quiz docx must
    end right after the last question/choice/ANSWER/POINT line — nothing else, since
    Quick Import keeps parsing anything left in the file."""
    fixture = Path(__file__).parent / "fixtures" / "short_answer_quiz.zip"
    result = convert_zip("short_answer_quiz.zip", fixture.read_bytes())
    assert result.success is True

    doc = Document(io.BytesIO(result.docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs]
    last_non_empty = next(t for t in reversed(paragraphs) if t.strip())
    assert last_non_empty == "16. 1001 1001"
    assert "Answer Key" not in "\n".join(paragraphs)


def test_not_a_zip_fails_gracefully():
    result = convert_zip("not_a_zip.zip", b"just some bytes, not a zip")
    assert result.success is False
    assert "not a valid ZIP" in result.error


def test_zip_with_no_html_fails_gracefully():
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("responses.csv", "a,b,c\n1,2,3\n")
    result = convert_zip("empty.zip", buf.getvalue())
    assert result.success is False
    assert result.error == "No Google Form HTML found"


def test_zip_slip_path_traversal_is_blocked():
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("../../evil.txt", "gotcha")
    result = convert_zip("evil.zip", buf.getvalue())
    assert result.success is False
    assert "unsafe" in result.error.lower() or "No Google Form HTML found" in result.error


def test_batch_isolation_one_bad_file_does_not_affect_another():
    good = FIXTURE_ZIP.read_bytes()
    bad = b"garbage"

    r1 = convert_zip("good.zip", good)
    r2 = convert_zip("bad.zip", bad)

    assert r1.success is True
    assert r2.success is False
